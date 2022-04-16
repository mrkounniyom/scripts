import docx

document = docx.Document('/home/mk/Downloads/tets.docx')

all_paras = document.paragraphs
print(len(all_paras))

for item in all_paras:
    print(item.text)

single = all_paras[0]

for run in single.runs:
    print(run)

mydoc = docx.Document()
mydoc.add_paragraph("This is first paragraph of a MS Word file.")
mydoc.save("my_written_file.docx")
